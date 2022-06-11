import argparse
import os
import pandas as pd
from docx import Document

def replace(document, mapping):
    for paragraph in document.paragraphs:
        for input_word, output_word in mapping.items():
            if input_word in paragraph.text:
                # input_format = paragraph.paragraph_format #'add_run', 'alignment', 'clear', 'insert_paragraph_before', 'paragraph_format', 'part', 'runs', 'style', 'text'
                paragraph_runs = paragraph.runs
                for i in range(len(paragraph_runs)):
                    input_text = paragraph_runs[i].text
                    if input_word in paragraph_runs[i].text:
                        output_text = input_text.replace(input_word, output_word)
                        paragraph_runs[i].text = output_text
                        print(" ", input_word, " -> ", output_word)
    return document


def main(template_path, data_path, output_dir):
    data = pd.read_csv(data_path, dtype=str)
    print(data_path)
    print(data)


    os.makedirs(output_dir, exist_ok=True)
    print(output_dir)

    for index, row in data.iterrows():
        document_template = Document(template_path)
        print(template_path)

        mapping = row.to_dict()
        print(index, mapping)

        document_result = replace(document_template, mapping)

        filename = row['FILE'] + ".docx"
        filename = filename.replace(" ", "_")
        output_file = os.path.join(output_dir, filename)
        document_result.save(output_file)
        print(index, output_file)
        print()


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", type=str, help="[input] template", default="template.docx")
    parser.add_argument("--data", type=str, help="[output] data", default="data.csv")
    parser.add_argument("--output", type=str, help="[output] destination", default="output/")
    args = parser.parse_args()
    main(args.template, args.data, args.output)